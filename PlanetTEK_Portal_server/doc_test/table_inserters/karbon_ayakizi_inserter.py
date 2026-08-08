import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def set_cell_background(cell, fill_hex):
    """Hücre arka plan rengini ayarlar."""
    if not fill_hex:
        return
    fill_hex = fill_hex.replace("#", "")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=30, bottom=30, left=40, right=40):
    """Hücre içi padding ayarları (Karbon ayakizi tablosu için optimize edildi)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old_mar in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old_mar)
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, **kwargs):
    """Hücre kenarlıklarını MS Word XML yapısına uygun olarak yazar."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, attr in [
                ("val", "w:val"),
                ("color", "w:color"),
                ("sz", "w:sz"),
                ("space", "w:space"),
            ]:
                if key in edge_data:
                    element.set(qn(attr), str(edge_data[key]))


def parse_inline_styles(style_str):
    """HTML inline CSS stringini dictionary olarak dönüştürür."""
    styles = {}
    if not style_str:
        return styles
    for item in style_str.split(";"):
        if ":" in item:
            k, v = item.split(":", 1)
            styles[k.strip().lower()] = v.strip()
    return styles


def parse_border_css(border_css_str):
    """CSS border metnini 1/4 pt (sz: 2) standart Word border objesine dönüştürür."""
    if not border_css_str:
        return None
    b_type = "double" if "double" in border_css_str else "single"
    return {"val": b_type, "sz": "2", "color": "000000", "space": "0"}


def html_to_word_karbon_ayakizi_table(html_content, doc):
    """Karbon Ayakizi HTML tablosunu dxa/tblGrid seviyesinde tam genişlikte Word'e aktarır."""
    soup = BeautifulSoup(html_content, "html.parser")
    html_table = soup.find("table")
    if not html_table:
        return None

    rows = html_table.find_all("tr")
    num_rows = len(rows)
    num_cols = 6  # Tablo tam 6 sütunlu

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = False
    table.allow_autofit = False

    # 6 Sütun için tam A4 genişliği (Toplam ~6.5 inç / 9360 dxa)
    col_widths_inches = [1.65, 0.70, 0.90, 1.65, 0.70, 0.90]
    col_widths_dxa = [int(w * 1440) for w in col_widths_inches]
    total_table_dxa = sum(col_widths_dxa)

    tblElem = table._element
    tblPr = tblElem.xpath("w:tblPr")[0]

    # 1. Tablo Genişliğini dxa (Sabit) Olarak Ayarla ve Autofit'i Kapat
    for old_w in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old_w)

    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_table_dxa}" w:type="dxa"/>')
    tblPr.append(tblW)

    tblAutofit = tblPr.find(qn("w:tblAutofit"))
    if tblAutofit is not None:
        tblAutofit.set(qn("w:val"), "0")
    else:
        tblPr.append(parse_xml(f'<w:tblAutofit {nsdecls("w")} w:val="0"/>'))

    # 2. <w:tblGrid> Yapısını XML Seviyesinde Oluştur (LibreOffice Sıkışmasını Engeller)
    old_grid = tblElem.find(qn("w:tblGrid"))
    if old_grid is not None:
        tblElem.remove(old_grid)

    tblGrid = OxmlElement("w:tblGrid")
    for w_dxa in col_widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w_dxa))
        tblGrid.append(gc)
    tblElem.insert(1, tblGrid)

    # 3. Hücre Yapılandırması ve İçerik Doldurma
    occupied = [[False] * num_cols for _ in range(num_rows)]

    for r_idx, tr in enumerate(rows):
        c_idx = 0
        cells = tr.find_all(["td", "th"])

        for td in cells:
            while c_idx < num_cols and occupied[r_idx][c_idx]:
                c_idx += 1

            if c_idx >= num_cols:
                break

            rowspan = int(td.get("rowspan", 1))
            colspan = int(td.get("colspan", 1))

            cell_start = table.cell(r_idx, c_idx)
            cell_end = table.cell(
                r_idx + rowspan - 1, min(num_cols - 1, c_idx + colspan - 1)
            )

            if rowspan > 1 or colspan > 1:
                merged_cell = cell_start.merge(cell_end)
            else:
                merged_cell = cell_start

            for r in range(r_idx, min(num_rows, r_idx + rowspan)):
                for c in range(c_idx, min(num_cols, c_idx + colspan)):
                    occupied[r][c] = True

            # Birleştirilmiş hücrelerin toplam dxa genişliğini hesapla
            cell_dxa_width = sum(col_widths_dxa[c_idx : min(num_cols, c_idx + colspan)])

            # XML seviyesinde w:tcW genişliği ataması
            tcPr = merged_cell._tc.get_or_add_tcPr()
            for old_tcW in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old_tcW)
            tcPr.append(
                parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{cell_dxa_width}" w:type="dxa"/>'
                )
            )

            # Stil ve Metin Biçimlendirmeleri
            style = parse_inline_styles(td.get("style", ""))

            bg_color = style.get("background-color")
            if bg_color:
                set_cell_background(merged_cell, bg_color)

            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = merged_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

            align = style.get("text-align", "left")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            text = td.get_text(separator="\n").strip()
            run = p.add_run(text)
            run.font.name = "Calibri"

            font_size_str = style.get("font-size", "8pt")
            size_match = re.search(r"[\d.]+", font_size_str)
            size_val = float(size_match.group()) if size_match else 8.0
            run.font.size = Pt(size_val)

            if style.get("font-weight") == "bold":
                run.bold = True
            if style.get("font-style") == "italic":
                run.italic = True

            color_str = style.get("color")
            if color_str and color_str.startswith("#"):
                hex_c = color_str.replace("#", "")
                run.font.color.rgb = RGBColor(
                    int(hex_c[0:2], 16),
                    int(hex_c[2:4], 16),
                    int(hex_c[4:6], 16),
                )

            border_opts = {}
            for edge in ["top", "bottom", "left", "right"]:
                b_key = f"border-{edge}"
                if b_key in style:
                    border_opts[edge] = parse_border_css(style[b_key])

            set_cell_border(merged_cell, **border_opts)
            set_cell_margins(merged_cell, top=30, bottom=30, left=40, right=40)

            c_idx += colspan

    return table


def process_karbon_ayakizi_and_insert():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_test_dir = os.path.abspath(os.path.join(script_dir, ".."))
    documents_dir = os.path.join(doc_test_dir, "Documents")
    tables_dir = os.path.join(doc_test_dir, "table_creators", "tables")

    raw_docx_path = os.path.join(documents_dir, "document_updated.docx")
    updated_docx_path = os.path.join(documents_dir, "document_updated.docx")

    target_docx_path = (
        updated_docx_path if os.path.exists(updated_docx_path) else raw_docx_path
    )
    output_path = updated_docx_path
    html_file_path = os.path.join(tables_dir, "htmls", "karbon_ayakizi.html")

    if not os.path.exists(target_docx_path):
        print(f"Hata: '{target_docx_path}' Word dosyası bulunamadı!")
        return

    if not os.path.exists(html_file_path):
        alt_html_path = os.path.join(script_dir, "karbon_ayakizi.html")
        if os.path.exists(alt_html_path):
            html_file_path = alt_html_path
        else:
            print(f"Hata: 'karbon_ayakizi.html' dosyası bulunamadı!")
            return

    print(f"📄 Word belgesi yükleniyor: {os.path.basename(target_docx_path)}")
    doc = Document(target_docx_path)

    placeholder = "##karbon_ayakizi##"

    with open(html_file_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text.lower():
            target_p = p
            break

    if target_p:
        print(
            f"  🔍 '{placeholder}' etiketi bulundu. Karbon Ayakizi tablosu ekleniyor..."
        )

        word_table = html_to_word_karbon_ayakizi_table(html_content, doc)

        if word_table:
            target_p._p.getparent().insert(
                target_p._p.getparent().index(target_p._p), word_table._tbl
            )

            p_element = target_p._p
            p_element.getparent().remove(p_element)

            print("  ✅ Karbon Ayakizi tablosu yerleştirildi ve etiket kaldırıldı.")
    else:
        print(
            f"  ⚠️ Belge içinde '{placeholder}' etiketi bulunamadı, tablo sona ekleniyor..."
        )
        html_to_word_karbon_ayakizi_table(html_content, doc)

    doc.save(output_path)
    print(f"\n🎉 Karbon Ayakizi Tablosu Aktarımı Tamamlandı! Dosya: {output_path}")


if __name__ == "__main__":
    process_karbon_ayakizi_and_insert()
