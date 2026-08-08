import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt


def set_cell_background(cell, fill_hex):
    """Hücre arka plan rengini ayarlar."""
    if not fill_hex:
        return
    fill_hex = fill_hex.replace("#", "")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=15, bottom=15, left=70, right=70):
    """Hücre içi padding ayarları (OPEX için dikey yüksekliği korur)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old_mar in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old_mar)
    tcMar = OxmlElement("w:tcMar")
    for m, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, **kwargs):
    """Hücre kenarlıklarını MS Word XML yapısına uygun olarak yazar."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, attr in [
                ("val", "w:val"),
                ("color", "w:color"),
                ("sz", "w:sz"),
                ("space", "w:space"),
            ]:
                if key in edge_data:
                    element.set(qn(attr), str(edge_data[key]))


def parse_inline_styles(style_str):
    """HTML inline CSS stringini dictionary olarak dönüştürür."""
    styles = {}
    if not style_str:
        return styles
    for item in style_str.split(";"):
        if ":" in item:
            k, v = item.split(":", 1)
            styles[k.strip().lower()] = v.strip()
    return styles


def parse_border_css(border_css_str):
    """CSS border metnini 1/4 pt (sz: 2) standart Word border objesine dönüştürür."""
    if not border_css_str:
        return None
    b_type = "double" if "double" in border_css_str else "single"
    return {"val": b_type, "sz": "2", "color": "000000"}


def html_to_word_opex_table(html_content, doc):
    """OPEX HTML tablosunu dxa/tblGrid seviyesinde 2 sütunlu yapıda oluşturur ve LibreOffice büzüşmesini engeller."""
    soup = BeautifulSoup(html_content, "html.parser")
    html_table = soup.find("table")
    if not html_table:
        return None

    rows = html_table.find_all("tr")
    num_rows = len(rows)
    num_cols = 2

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # OPEX Tablosu Özel A4 Sütun Genişlikleri (4.5 inç + 2.0 inç = 6.5 inç / 9360 dxa)
    col_widths_inches = [4.50, 2.00]
    col_widths_dxa = [int(w * 1440) for w in col_widths_inches]
    total_table_dxa = sum(col_widths_dxa)

    tblElem = table._element
    tblPr = tblElem.xpath("w:tblPr")[0]

    # 1. Tablo Genişliğini dxa (Sabit) Olarak Ayarla ve Autofit'i Kapat
    for old_w in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old_w)

    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_table_dxa}" w:type="dxa"/>')
    tblPr.append(tblW)

    tblAutofit = tblPr.find(qn("w:tblAutofit"))
    if tblAutofit is not None:
        tblAutofit.set(qn("w:val"), "0")
    else:
        tblPr.append(parse_xml(f'<w:tblAutofit {nsdecls("w")} w:val="0"/>'))

    # Tablo Düzenini Kesin Olarak Sabit (Fixed) Yap
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is not None:
        tblLayout.set(qn("w:type"), "fixed")
    else:
        tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

    # 2. <w:tblGrid> Yapısını XML Seviyesinde Oluştur (LibreOffice Sıkışmasını Engeller)
    old_grid = tblElem.find(qn("w:tblGrid"))
    if old_grid is not None:
        tblElem.remove(old_grid)

    tblGrid = OxmlElement("w:tblGrid")
    for w_dxa in col_widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w_dxa))
        tblGrid.append(gc)
    tblElem.insert(1, tblGrid)

    # 3. Hücre Yapılandırması ve İçerik Doldurma
    occupied = [[False] * num_cols for _ in range(num_rows)]

    for r_idx, tr in enumerate(rows):
        c_idx = 0
        cells = tr.find_all(["td", "th"])

        # Satırların Sayfa Kırılımında Bölünmesini Engelle
        trPr = table.rows[r_idx]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f"<w:cantSplit {nsdecls('w')}/>"))

        for td in cells:
            while c_idx < num_cols and occupied[r_idx][c_idx]:
                c_idx += 1

            if c_idx >= num_cols:
                break

            rowspan = int(td.get("rowspan", 1))
            colspan = int(td.get("colspan", 1))

            cell_start = table.cell(r_idx, c_idx)
            cell_end = table.cell(
                r_idx + rowspan - 1, min(num_cols - 1, c_idx + colspan - 1)
            )

            if rowspan > 1 or colspan > 1:
                merged_cell = cell_start.merge(cell_end)
            else:
                merged_cell = cell_start

            for r in range(r_idx, min(num_rows, r_idx + rowspan)):
                for c in range(c_idx, min(num_cols, c_idx + colspan)):
                    occupied[r][c] = True

            # Merged (colspan) hücrelerde her alt grid hücresine kendi dxa genişliğini yaz
            for c_offset in range(colspan):
                curr_c = c_idx + c_offset
                if curr_c < num_cols:
                    target_cell = table.cell(r_idx, curr_c)
                    tcPr = target_cell._tc.get_or_add_tcPr()
                    for old_tcW in tcPr.findall(qn("w:tcW")):
                        tcPr.remove(old_tcW)
                    tcPr.append(
                        parse_xml(
                            f'<w:tcW {nsdecls("w")} w:w="{col_widths_dxa[curr_c]}"'
                            ' w:type="dxa"/>'
                        )
                    )

            # Stil ve Metin Biçimlendirmeleri
            style = parse_inline_styles(td.get("style", ""))

            bg_color = style.get("background-color")
            if bg_color:
                set_cell_background(merged_cell, bg_color)

            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = merged_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

            align = style.get("text-align", "left")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            text = td.get_text(separator="\n").strip()
            run = p.add_run(text)
            run.font.name = "Calibri"

            font_size_str = style.get("font-size", "8pt")
            size_match = re.search(r"[\d.]+", font_size_str)
            size_val = float(size_match.group()) if size_match else 9.0
            run.font.size = Pt(size_val)

            if style.get("font-weight") == "bold":
                run.bold = True
            if style.get("font-style") == "italic":
                run.italic = True

            border_opts = {}
            for border_name in ["top", "bottom", "left", "right"]:
                css_border_key = f"border-{border_name}"
                if css_border_key in style:
                    b_str = style[css_border_key]
                    b_type = "double" if "double" in b_str else "single"
                    border_opts[border_name] = {
                        "val": b_type,
                        "sz": "2",  # 1/4 pt (0.25 pt)
                        "color": "000000",
                    }

            set_cell_border(merged_cell, **border_opts)
            set_cell_margins(merged_cell, top=15, bottom=15, left=70, right=70)

            c_idx += colspan

    return table


def process_opex_and_insert():
    """Yeni Klasör Yapısına Uyumlu OPEX Aktarım Süreci."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_test_dir = os.path.abspath(os.path.join(script_dir, ".."))
    root_dir = os.path.abspath(os.path.join(script_dir, ".."))

    tables_dir = os.path.join(doc_test_dir, "table_creators", "tables")
    documents_dir = os.path.join(root_dir, "Documents")

    docx_path = os.path.join(documents_dir, "document_updated.docx")
    output_path = os.path.join(documents_dir, "document_updated.docx")
    opex_html_path = os.path.join(tables_dir, "htmls", "opex.html")

    if not os.path.exists(docx_path):
        print(f"Hata: '{docx_path}' Word dosyası bulunamadı!")
        return

    if not os.path.exists(opex_html_path):
        print(f"Hata: '{opex_html_path}' bulunamadı!")
        return

    print(f"📄 Word belgesi yükleniyor: {os.path.basename(docx_path)}")
    doc = Document(docx_path)

    placeholder = "##opex##"

    with open(opex_html_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text.lower():
            target_p = p
            break

    if target_p:
        print(f"  🔍 '{placeholder}' etiketi bulundu. OPEX tablosu işleniyor...")

        word_table = html_to_word_opex_table(html_content, doc)

        if word_table:
            target_p._p.getparent().insert(
                target_p._p.getparent().index(target_p._p), word_table._tbl
            )

            p_element = target_p._p
            p_element.getparent().remove(p_element)

            print("  ✅ OPEX tablosu yerleştirildi ve etiket paragrafı temizlendi.")
    else:
        print(
            f"  ⚠️ Belge içinde '{placeholder}' etiketi bulunamadı, tablo sona"
            " ekleniyor..."
        )
        html_to_word_opex_table(html_content, doc)

    doc.save(output_path)
    print(f"\n🎉 OPEX Aktarımı Tamamlandı! Dosya: {output_path}")


if __name__ == "__main__":
    process_opex_and_insert()
