import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def set_cell_background(cell, fill_hex):
    """Hücre arka plan rengini ayarlar."""
    if not fill_hex:
        return
    fill_hex = fill_hex.replace("#", "")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    """Kapak tablosu dikey yükseklik ve iç dolguları (padding)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old_mar in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old_mar)
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, **kwargs):
    """Hücre kenarlıklarını MS Word XML yapısına uygun olarak yazar."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, attr in [
                ("val", "w:val"),
                ("color", "w:color"),
                ("sz", "w:sz"),
                ("space", "w:space"),
            ]:
                if key in edge_data:
                    element.set(qn(attr), str(edge_data[key]))


def apply_table_level_outer_double_borders(table):
    """Kapak tablosu genelinde dış çift çerçeveyi garantiye alır."""
    tblPr = table._tbl.tblPr
    for old_borders in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old_borders)

    tblBorders = OxmlElement("w:tblBorders")
    borders_dict = {
        "top": {"val": "double", "sz": "4", "space": "0", "color": "000000"},
        "left": {"val": "double", "sz": "4", "space": "0", "color": "000000"},
        "bottom": {"val": "double", "sz": "4", "space": "0", "color": "000000"},
        "right": {"val": "double", "sz": "4", "space": "0", "color": "000000"},
    }

    for b_name, b_attrs in borders_dict.items():
        border_el = OxmlElement(f"w:{b_name}")
        for k, v in b_attrs.items():
            border_el.set(qn(f"w:{k}"), v)
        tblBorders.append(border_el)

    tblPr.append(tblBorders)


def parse_inline_styles(style_str):
    """HTML inline CSS stringini dictionary olarak dönüştürür."""
    styles = {}
    if not style_str:
        return styles
    for item in style_str.split(";"):
        if ":" in item:
            k, v = item.split(":", 1)
            styles[k.strip().lower()] = v.strip()
    return styles


def parse_border_css(border_css_str):
    """CSS border metnini Word border objesine dönüştürür."""
    if not border_css_str:
        return None
    b_type = "double" if "double" in border_css_str else "single"
    return {"val": b_type, "sz": "4", "color": "000000"}


def html_to_word_kapak_table(html_content, doc):
    """Kapak HTML tablosunu 2 sütunlu yapıda oluşturur ve sayfada ortalar."""
    soup = BeautifulSoup(html_content, "html.parser")
    html_table = soup.find("table")
    if not html_table:
        return None

    rows = html_table.find_all("tr")
    num_rows = len(rows)
    num_cols = 2  # Sadece A ve B Kolonları

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = False
    table.allow_autofit = False

    # TABLOYU SAYFADA YATAY ORTALAMA
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Kapak Tablosu Kolon Genişlikleri (Sol Etiket: ~3.8 inç, Sağ Değer: ~2.2 inç | Toplam: ~6.0 inç)
    col_widths = [Inches(3.8), Inches(2.2)]

    tblPr = table._tbl.tblPr
    for old_w in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(col_widths) * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    occupied = [[False] * num_cols for _ in range(num_rows)]

    for r_idx, tr in enumerate(rows):
        c_idx = 0
        cells = tr.find_all(["td", "th"])

        for td in cells:
            while c_idx < num_cols and occupied[r_idx][c_idx]:
                c_idx += 1

            if c_idx >= num_cols:
                break

            rowspan = int(td.get("rowspan", 1))
            colspan = int(td.get("colspan", 1))

            cell_start = table.cell(r_idx, c_idx)
            cell_end = table.cell(
                r_idx + rowspan - 1, min(num_cols - 1, c_idx + colspan - 1)
            )

            if rowspan > 1 or colspan > 1:
                merged_cell = cell_start.merge(cell_end)
            else:
                merged_cell = cell_start

            for r in range(r_idx, min(num_rows, r_idx + rowspan)):
                for c in range(c_idx, min(num_cols, c_idx + colspan)):
                    occupied[r][c] = True

            style = parse_inline_styles(td.get("style", ""))

            bg_color = style.get("background-color")
            if bg_color:
                set_cell_background(merged_cell, bg_color)

            # Hücre içi Paragraf Biçimlendirmesi
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = merged_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05

            align = style.get("text-align", "left")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            text = td.get_text(separator="\n").strip()
            run = p.add_run(text)
            run.font.name = "Calibri"

            font_size_str = style.get("font-size", "9pt")
            size_match = re.search(r"[\d.]+", font_size_str)
            size_val = float(size_match.group()) if size_match else 9.0
            run.font.size = Pt(size_val)

            if style.get("font-weight") == "bold":
                run.bold = True
            if style.get("font-style") == "italic":
                run.italic = True

            # CSS'ten Gelen Kenarlıklar
            border_opts = {}
            if "border-top" in style:
                border_opts["top"] = parse_border_css(style["border-top"])
            if "border-bottom" in style:
                border_opts["bottom"] = parse_border_css(style["border-bottom"])
            if "border-left" in style:
                border_opts["left"] = parse_border_css(style["border-left"])
            if "border-right" in style:
                border_opts["right"] = parse_border_css(style["border-right"])

            set_cell_border(merged_cell, **border_opts)
            set_cell_margins(merged_cell, top=60, bottom=60, left=100, right=100)

            c_idx += colspan

    # Dış Çerçeveyi Çift Çizgi Yap
    apply_table_level_outer_double_borders(table)

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    return table


def process_kapak_and_insert():
    # table_inserters -> doc_test
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_test_dir = os.path.abspath(os.path.join(script_dir, ".."))

    documents_dir = os.path.join(doc_test_dir, "Documents")
    tables_dir = os.path.join(doc_test_dir, "table_creators", "tables")

    # ÖNEMLİ: Dosya Okuma Sıralaması
    # Eğer ham şablon (document.docx) varsa ve etiketi arıyorsak, target her zaman document.docx olmalı
    # veya güncellenmiş dosya varsa ondan devam etmeli.
    raw_docx_path = os.path.join(documents_dir, "document.docx")
    updated_docx_path = os.path.join(documents_dir, "document_updated.docx")

    target_docx_path = (
        updated_docx_path if os.path.exists(updated_docx_path) else raw_docx_path
    )
    output_path = updated_docx_path
    kapak_html_path = os.path.join(tables_dir, "..", "tables", "htmls", "kapak.html")

    if not os.path.exists(target_docx_path):
        print(f"Hata: '{target_docx_path}' Word dosyası bulunamadı!")
        return

    if not os.path.exists(kapak_html_path):
        print(f"Hata: '{kapak_html_path}' bulunamadı!")
        return

    print(f"📄 Word belgesi yükleniyor: {os.path.basename(target_docx_path)}")
    doc = Document(target_docx_path)

    placeholder = "##kapak##"

    with open(kapak_html_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text.lower():
            target_p = p
            break

    if target_p:
        print(f"  🔍 '{placeholder}' etiketi bulundu. Kapak tablosu ekleniyor...")

        word_table = html_to_word_kapak_table(html_content, doc)

        if word_table:
            target_p._p.getparent().insert(
                target_p._p.getparent().index(target_p._p), word_table._tbl
            )

            p_element = target_p._p
            p_element.getparent().remove(p_element)

            print("  ✅ Kapak tablosu tam yerleştirildi ve etiket kaldırıldı.")
    else:
        print(f"  ⚠️ Belge içinde '{placeholder}' etiketi bulunamadı!")

    doc.save(output_path)
    print(f"\n🎉 Kapak Tablosu Aktarımı Tamamlandı! Dosya: {output_path}")


if __name__ == "__main__":
    process_kapak_and_insert()
