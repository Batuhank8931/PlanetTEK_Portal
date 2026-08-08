import os
import json
import re
import datetime
from docx import Document


def get_formatted_date(language):
    """
    Bugünün tarihini teklif diline göre formatlar.
    Yerli -> 24 Nisan 2026
    Yabancı -> April 24, 2026
    """
    now = datetime.datetime.now()
    day = now.day

    tr_months = [
        "Ocak",
        "Şubat",
        "Mart",
        "Nisan",
        "Mayıs",
        "Haziran",
        "Temmuz",
        "Ağustos",
        "Eylül",
        "Ekim",
        "Kasım",
        "Aralık",
    ]
    en_months = [
        "January",
        "February",
        "March",
        "April",
        "May",
        "June",
        "July",
        "August",
        "September",
        "October",
        "November",
        "December",
    ]

    if language == "yerli":
        month = tr_months[now.month - 1]
        return f"{day} {month} {now.year}"
    else:
        month = en_months[now.month - 1]
        return f"{month} {day}, {now.year}"


def parse_offer_number(offer_number):
    """
    offer_number string'ini ayrıştırarak module_details1 ve module_details2 değerlerini üretir.
    Örnek offer_number: "YDS R0 06 08 2026 1 MX 1 70 25 114"
    """
    parts = offer_number.strip().split()
    if not parts:
        return "", ""

    # İlk parametre: ##module detailes2## (örn: YDS)
    module_details2 = parts[0]

    # Tarih kısmını tespit et (GG MM YYYY biçimindeki 3 blok)
    date_start_idx = -1
    for i in range(len(parts) - 2):
        if parts[i].isdigit() and parts[i + 1].isdigit() and parts[i + 2].isdigit():
            if len(parts[i + 2]) == 4:  # Yıl 4 haneli ise
                date_start_idx = i
                break

    module_details1 = ""
    if date_start_idx != -1:
        after_date_parts = parts[date_start_idx + 3 :]

        model_parts = []
        for p in after_date_parts:
            if p.isdigit() and len(p) >= 2 and int(p) > 10:
                break
            model_parts.append(p)

        module_details1 = " ".join(model_parts)

    return module_details1, module_details2


def get_capacity_string(form_data):
    """
    planetDiskDetails.debi değerini alır.
    unitSystem == 'US' / 'IMPERIAL' ise GPD'ye çevirir (1 m3/gün = 264.172 GPD).
    teklifDili 'yerli' ise m³/gün, 'yabancı' ise m³/day yazar.
    """
    customer_info = form_data.get("customerInfo", {})
    planet_disk = form_data.get("planetDiskDetails", {})

    unit_system = str(customer_info.get("unitSystem", "")).strip().upper()
    teklif_dili = str(customer_info.get("teklifDili", "")).strip().lower()

    # debi değerini sayısal türe çevirme
    raw_debi = planet_disk.get("debi", 0)
    try:
        debi_val = float(raw_debi)
    except (ValueError, TypeError):
        debi_val = 0.0

    # US / Imperial Birim Çevrimi
    if unit_system in ["US", "IMPERIAL"]:
        gpd_val = round(debi_val * 264.172)
        # 1.000 ayırıcı formatı (örn: 18.492 GPD)
        return f"{gpd_val:,}".replace(",", ".") + " GPD"
    else:
        # Metrik Birim Formatı
        # Tam sayıysa küsüratsız, küsüratlıysa düzgün string çevrimi
        formatted_debi = f"{int(debi_val)}" if debi_val.is_integer() else f"{debi_val}"
        unit_label = "m³/gün" if teklif_dili == "yerli" else "m³/day"
        return f"{formatted_debi} {unit_label}"


def replace_placeholders_in_doc(doc, replacements):
    """
    Word belgesindeki font ve stil biçimlendirmelerini bozmadan
    placeholder metinlerini değiştirir.
    """

    def replace_in_paragraph(paragraph):
        for key, val in replacements.items():
            if key in paragraph.text:
                # 1. Durum: Placeholder tek bir Run içerisindeyse (stili birebir korur)
                replaced = False
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(val))
                        replaced = True

                # 2. Durum: Word placeholder'ı birden fazla Run'a böldüyse
                if not replaced and key in paragraph.text:
                    # İlk run'ın stilini koruyarak tüm metni birleştirir
                    full_text = paragraph.text.replace(key, str(val))
                    paragraph.runs[0].text = full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""

    # Paragrafları tara
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # Tablo hücrelerini tara
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def prepare_final_document():
    # 1. Yollar
    base_dir = os.path.dirname(os.path.abspath(__file__))
    json_path = os.path.join(base_dir, "formData.json")
    docs_dir = os.path.join(base_dir, "Documents")
    source_docx = os.path.join(docs_dir, "table_arranged_document.docx")

    # 2. Kontroller
    if not os.path.exists(json_path):
        print(f"Hata: '{json_path}' bulunamadı!")
        return

    if not os.path.exists(source_docx):
        print(f"Hata: Kaynak döküman '{source_docx}' bulunamadı!")
        return

    # 3. JSON Okuma
    with open(json_path, "r", encoding="utf-8") as f:
        form_data = json.load(f)

    customer_info = form_data.get("customerInfo", {})
    teklif_dili_raw = customer_info.get("teklifDili", "").strip()
    teklif_dili = teklif_dili_raw.lower()

    # 4. Değişkenler
    company_name = customer_info.get("ticari_unvan", "")
    customer_email = customer_info.get("ilgiliKisi_email", "")
    customer_name = customer_info.get("ilgiliKisi", "")

    # Dinamik Kapasite (Debi) Hesabı
    compacity = get_capacity_string(form_data)

    formatted_date = get_formatted_date(teklif_dili)
    date_today_short = datetime.datetime.now().strftime("%d.%m.%Y")

    # Offer number parsing
    offer_number = customer_info.get("offer_number", "")
    module_details1, module_details2 = parse_offer_number(offer_number)

    # Replace haritası
    replacements = {
        "##company_name##": company_name,
        "##compacity##": compacity,
        "##date##": formatted_date,
        "##customer_email##": customer_email,
        "##datel##": formatted_date,
        "##customer_name##": customer_name,
        "##delivery1##": "1",
        "##delivery2##": "8-10",
        "##delivery3##": "3-4",
        "##delivery4##": "1",
    }

    # 5. Döküman İşleme
    doc = Document(source_docx)
    replace_placeholders_in_doc(doc, replacements)

    # 6. Dosya Adı Oluşturma
    if teklif_dili == "yerli":
        output_filename = f"PlanetTEK Teklif {company_name} - {module_details1} - {module_details2} - {date_today_short} - {compacity}.docx"
    else:
        output_filename = f"PlanetTEK Offer {company_name} - {module_details1} - {module_details2} - {date_today_short} - {compacity}.docx"

    # Geçersiz karakter temizliği (örn: / \ : * ? " < > |)
    output_filename = re.sub(r'[\\/*?:"<>|]', "", output_filename)
    output_path = os.path.join(docs_dir, output_filename)

    # 7. Kaydetme
    try:
        doc.save(output_path)
        print("İşlem başarıyla tamamlandı!")
        print(f"Hesaplanan Kapasite: {compacity}")
        print(f"Yeni Dosya Oluşturuldu: {output_path}")
    except Exception as e:
        print(f"Dosya kaydedilirken hata oluştu: {e}")


if __name__ == "__main__":
    prepare_final_document()
