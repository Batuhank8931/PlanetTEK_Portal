import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Word hücresine mutlak dolgu rengini basar."""
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    for old_shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(old_shd)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shading_elm)

def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    """Hücre içi boşlukları (padding) ayarlar."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old_mar in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old_mar)
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top="single", bottom="single", left="single", right="single", color="1C1C1C", sz="4"):
    """Kenarlıkları Word yerel grid çizgilerini ezerek enjekte eder."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old_borders in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old_borders)
        
    tcBorders = OxmlElement('w:tcBorders')
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    
    for border_name, border_style in borders.items():
        if border_style:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), border_style)
            node.set(qn('w:sz'), "12" if border_style == "double" else sz)
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), color)
            tcBorders.append(node)
            
    tcPr.append(tcBorders)

def insert_html_table_to_docx(html_path="table.html", docx_path="document.docx", output_path="document_updated.docx"):
    if not os.path.exists(html_path):
        print(f"Hata: '{html_path}' bulunamadı!")
        return
    if not os.path.exists(docx_path):
        print(f"Bilgi: '{docx_path}' bulunamadı, şablon üretiliyor...")
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("##TABLO_ALANI##")
        doc.save(docx_path)

    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    
    html_table = soup.find("table")
    if not html_table:
        print("Hata: HTML tablosu bulunamadı!")
        return

    html_rows = html_table.find_all("tr")
    
    # Sabit daraltılmış kolon haritamız
    col_widths_px = [55, 275, 40, 70, 75, 65, 90]
    max_cols = len(col_widths_px)

    doc = Document(docx_path)
    target_p = None
    for p in doc.paragraphs:
        if "##TABLO_ALANI##" in p.text:
            target_p = p
            break
            
    if not target_p:
        print("Hata: '##TABLO_ALANI##' bulunamadı!")
        return

    target_p.text = target_p.text.replace("##TABLO_ALANI##", "")

    # A4 Oranlama Havuzu (6.1 İnç)
    total_px = sum(col_widths_px)
    target_total_inches = 6.5  
    col_widths_inches = [(w_px / total_px) * target_total_inches for w_px in col_widths_px]

    print(f"-> Güvenli OpenXML yapısıyla tablo inşa ediliyor...")
    
    word_table = doc.add_table(rows=len(html_rows), cols=max_cols)
    word_table.autofit = False
    word_table.allow_autofit = False
    word_table.style = 'Normal Table'
    
    # 🛠️ GÜVENLİ GRID ENJEKSİYONU (Word Dosya Bozulma Hatasını Çözen Alan)
    # Mevcut tblGrid yapılarını süpürüyoruz
    for old_tblGrid in word_table._tbl.findall(qn('w:tblGrid')):
        word_table._tbl.remove(old_tblGrid)
        
    new_tblGrid = OxmlElement('w:tblGrid')
    for w_inch in col_widths_inches:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w_inch * 1440)))
        new_tblGrid.append(gridCol)
        
    # XML Şemasına göre tblGrid, tblPr elemanından HEMEN SONRA gelmelidir. Satırlardan sonraya eklenirse dosya bozulur.
    word_table._tbl.insert(1, new_tblGrid)

    # Global borders temizliği
    tblPr = word_table._tbl.tblPr
    for old_tbl_borders in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old_tbl_borders)

    visited = [[0 for _ in range(max_cols)] for _ in range(len(html_rows))]

    for r_idx, html_row in enumerate(html_rows):
        html_cells = html_row.find_all(["td", "th"])
        c_idx = 0
        
        for html_cell in html_cells:
            while c_idx < max_cols and visited[r_idx][c_idx] == 1:
                c_idx += 1
                
            if c_idx >= max_cols:
                break
                
            colspan = int(html_cell.get("colspan", 1))
            rowspan = int(html_cell.get("rowspan", 1))
            
            main_cell = word_table.cell(r_idx, c_idx)
            
            if colspan > 1 or rowspan > 1:
                target_r = r_idx + rowspan - 1
                target_c = c_idx + colspan - 1
                merge_to_cell = word_table.cell(target_r, target_c)
                main_cell.merge(merge_to_cell)
                
                for r_m in range(r_idx, target_r + 1):
                    for c_m in range(c_idx, target_c + 1):
                        visited[r_m][c_m] = 1
            else:
                visited[r_idx][c_idx] = 1

            # Hücre içi genişlik kilidi (w:tcW)
            cell_width_inch = sum(col_widths_inches[c_idx + i] for i in range(colspan) if (c_idx + i) < len(col_widths_inches))
            tcPr = main_cell._tc.get_or_add_tcPr()
            for old_w in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old_w)
            
            new_w = OxmlElement('w:tcW')
            new_w.set(qn('w:w'), str(int(cell_width_inch * 1440)))
            new_w.set(qn('w:type'), 'dxa')
            tcPr.append(new_w)
            main_cell.width = Inches(cell_width_inch)

            style_str = html_cell.get("style", "")
            
            bg_match = re.search(r'background-color:\s*#([0-9A-Fa-f]{6})', style_str)
            if bg_match:
                set_cell_background(main_cell, bg_match.group(1).upper())
            else:
                set_cell_background(main_cell, "FFFFFF")

            b_top = "double" if "border-top: 3px double" in style_str else "single"
            b_bottom = "double" if "border-bottom: 3px double" in style_str else "single"
            
            set_cell_borders(main_cell, top=b_top, bottom=b_bottom, left="single", right="single", color="1C1C1C", sz="4")
            set_cell_margins(main_cell, top=40, bottom=40, left=60, right=60)

            # Yazı İçeriği
            main_cell.text = html_cell.get_text().strip()
            p_docx = main_cell.paragraphs[0]
            
            p_format = p_docx.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            p_format.line_spacing = 1.0

            if "text-align: center" in style_str:
                p_docx.alignment = 1
            elif "text-align: right" in style_str:
                p_docx.alignment = 2
            else:
                p_docx.alignment = 0

            for run in p_docx.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(7.5)
                if "font-weight: bold" in style_str:
                    run.bold = True
                if "font-style: italic" in style_str:
                    run.italic = True
                    
                color_match = re.search(r'color:\s*#([0-9A-Fa-f]{6})', style_str)
                if color_match:
                    c_hex = color_match.group(1)
                    run.font.color.rgb = RGBColor(int(c_hex[0:2], 16), int(c_hex[2:4], 16), int(c_hex[4:6], 16))

            c_idx += colspan

    for idx, width_inch in enumerate(col_widths_inches):
        if idx < len(word_table.columns):
            word_table.columns[idx].width = Inches(width_inch)

    target_p._p.getparent().insert(target_p._p.getparent().index(target_p._p) + 1, word_table._tbl)
    doc.save(output_path)
    print(f"\n🎉 BOZULMA HATASI ÇÖZÜLDÜ! Dosya güvenle '{output_path}' olarak kaydedildi.")

if __name__ == "__main__":
    insert_html_table_to_docx()